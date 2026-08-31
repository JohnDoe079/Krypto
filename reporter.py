"""Generuje raport DOCX z wyników analizy raportów giełdowych."""

import os
import re
from datetime import datetime
from typing import Dict, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from models.schemas import ExtractedIdentifiers, AssetTransaction, _normalize_decimal, is_valid_ip
from matcher import ReportComparator

SECTION_TRANSLATIONS = {
    "Basic Information": "Basic Information (Podstawowe informacje)",
    "API Information": "API Information (Informacje API)",
    "KYC Approved Info": "KYC Approved Info (Zatwierdzone dane KYC)",
    "KYC Info": "KYC Info (Informacje KYC)",
    "Address Validation": "Address Validation (Weryfikacja adresu)",
    "EDD Info": "EDD Info (Informacje EDD – Rozszerzona due diligence)",
    "Payment Merchant Info": "Payment Merchant Info (Informacje o sprzedawcy)",
    "Sub-accounts": "Sub-accounts (Subkonta)",
    "HTX Register": "HTX Register (Rejestracja HTX)",
}

BASIC_INFO_ORDER = [
    "Registration time", "User ID", "User authentication type",
    "User authentication time", "User nationality", "User ID number",
    "First Name", "Last Name", "Mobile", "Email", "Status",
    "2FA", "2FA Opening time", "2FA Reset Information",
    "SMS", "SMS Opening time", "Tax ID", "TnC", "TnC Sign Date",
]

PAGE_WIDTH_INCHES = 7.0


def _to_float(val: str) -> float:
    """Konwertuje string na float, obsługuje przecinek dziesiętny."""
    try:
        normalized = _normalize_decimal(val)
        return float(normalized)
    except (ValueError, TypeError):
        return 0.0


def _is_pending_transaction(t: AssetTransaction) -> bool:
    """Sprawdza czy transakcja jest pending/niepotwierdzona (pomijamy w podsumowaniu)."""
    reason = str(t.reason).lower().strip() if t.reason else ""
    if not reason:
        return False

    ok_statuses = ["completed", "success", "filled", "confirmed"]
    if any(s in reason for s in ok_statuses):
        return False

    pending_kw = ["pending", "processing", "initiated", "created", "request", "order"]
    has_pending = any(kw in reason for kw in pending_kw)
    has_success = "success" in reason or "completed" in reason
    if has_pending and not has_success:
        return True

    success_required = ["withdrawal", "deposit", "transfer"]
    for op in success_required:
        if op in reason and not has_success:
            if any(x in reason for x in ["fee", "commission", "reward", "interest", "staking", "airdrop"]):
                return False
            return True

    return False


def _sum_asset_flow(transactions: List[AssetTransaction]) -> Dict[str, Dict[str, float]]:
    """Podsumowanie przepływów per waluta: przychody, rozchody, netto."""
    result = {}
    skipped = 0
    for t in transactions:
        if _is_pending_transaction(t):
            skipped += 1
            continue
        curr = t.currency.upper() if t.currency else "UNKNOWN"
        if curr not in result:
            result[curr] = {"in": 0.0, "out": 0.0, "count_in": 0, "count_out": 0}
        chg = _to_float(t.change)
        if chg > 1e-12:
            result[curr]["in"] += chg
            result[curr]["count_in"] += 1
        elif chg < -1e-12:
            result[curr]["out"] += abs(chg)
            result[curr]["count_out"] += 1

    if skipped > 0:
        print(f"  Pominięto {skipped} transakcji pending/niepotwierdzonych")

    return {k: v for k, v in result.items() if v["in"] > 1e-12 or v["out"] > 1e-12}


def _fmt(v: float) -> str:
    if abs(v) < 1e-12:
        return "0"
    s = f"{abs(v):.12f}".rstrip("0").rstrip(".")
    if s.startswith("."):
        s = "0" + s
    return s


def _fmt_signed(v: float) -> str:
    if abs(v) < 1e-12:
        return "0"
    s = _fmt(v)
    return f"+{s}" if v > 0 else f"-{s}"


def _is_fiat_duplicate(t: AssetTransaction, all_txns: List[AssetTransaction]) -> bool:
    """Sprawdza czy transakcja fiat jest duplikatem w innych logach (Spot/Funding/Deposit/Withdrawal).

    Kryteria duplikatu:
    - Ta sama waluta
    - Ta sama kwota (±0.1%)
    - Czas w zakresie ±5 minut
    - Źródło NIE jest Fiat (czyli Spot/Funding/Deposit/Withdrawal)
    """
    if t.wallet_type != "Fiat":
        return False

    t_curr = t.currency.upper() if t.currency else ""
    t_chg = _to_float(t.change)
    t_time = t.time[:19] if t.time else ""
    if not t_curr or abs(t_chg) < 1e-12 or not t_time:
        return False

    tolerance = max(1e-8, abs(t_chg) * 0.001)  # 0.1% lub minimum 1e-8

    for other in all_txns:
        if other.wallet_type == "Fiat" or other is t:
            continue
        o_curr = other.currency.upper() if other.currency else ""
        if o_curr != t_curr:
            continue
        o_chg = _to_float(other.change)
        if abs(o_chg - t_chg) > tolerance:
            continue
        # Sprawdź czas ±5 minut
        o_time = other.time[:19] if other.time else ""
        if not o_time:
            continue
        try:
            from datetime import datetime, timedelta
            t_dt = datetime.strptime(t_time, "%Y-%m-%d %H:%M:%S")
            o_dt = datetime.strptime(o_time, "%Y-%m-%d %H:%M:%S")
            if abs((t_dt - o_dt).total_seconds()) <= 300:  # 5 minut
                return True
        except ValueError:
            # Jeśli nie da się sparsować, porównaj stringi (dokładność do sekundy)
            if t_time == o_time:
                return True

    return False


class ReportGenerator:
    def __init__(self, output_path: str = "Raport_Analiza.docx"):
        self.output_path = output_path
        self.doc = Document()
        self._setup_styles()
        self._setup_margins()

    def _setup_margins(self):
        section = self.doc.sections[0]
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _add_heading(self, text: str, level: int = 1):
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if level == 1:
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
                run.font.size = Pt(20)
        elif level == 2:
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x00, 0x40, 0x80)
                run.font.size = Pt(14)
        elif level == 3:
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x00, 0x60, 0x80)
                run.font.size = Pt(12)
        elif level == 4:
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                run.font.size = Pt(11)
                run.bold = True
        return heading

    def _add_paragraph(self, text: str, bold: bool = False, color: RGBColor = None):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        return p

    def _add_table(self, headers: List[str], rows: List[List[str]], max_rows: int = None, col_widths: List[Inches] = None):
        if not rows:
            self._add_paragraph("(brak danych)")
            return None

        n_cols = len(headers)
        table = self.doc.add_table(rows=1, cols=n_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.allow_autofit = False

        for i in range(n_cols):
            if col_widths and i < len(col_widths):
                table.columns[i].width = col_widths[i]
            else:
                table.columns[i].width = Inches(PAGE_WIDTH_INCHES / n_cols)

        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            shading_elm = parse_xml(r'<w:shd {} w:fill="1F4E79"/>'.format(nsdecls('w')))
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        limit = max_rows if max_rows is not None else len(rows)
        for row_data in rows[:limit]:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                txt = str(cell_text) if cell_text is not None else ""
                row_cells[i].text = txt[:2000]
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        if max_rows is not None and len(rows) > max_rows:
            self._add_paragraph(f" ... i {len(rows) - max_rows} więcej wierszy")
        return table

    def _sort_basic_info(self, data: Dict[str, str]) -> List[List[str]]:
        order_map = {name: idx for idx, name in enumerate(BASIC_INFO_ORDER)}
        items = list(data.items())
        priority = []
        rest = []
        for col, val in items:
            if col in order_map:
                priority.append((order_map[col], col, val))
            else:
                rest.append((col, val))
        priority.sort(key=lambda x: x[0])
        rows = [[col, val if val is not None and str(val).strip() != "" else "(puste)"]
                for _, col, val in priority]
        rows += [[col, val if val is not None and str(val).strip() != "" else "(puste)"]
                 for col, val in rest]
        return rows

    # ========================================================================
    # HTX: NOWA STRUKTURA RAPORTU — PROFIL UŻYTKOWNIKA
    # ========================================================================

    def _render_htx_profile(self, r: ExtractedIdentifiers):
        """Renderuje pełny profil HTX: dane użytkownika → zdjęcia → portfele → transakcje."""
        htx_data = r.customer_info_sections.get("HTX Register", {})
        if not htx_data:
            self._add_paragraph("Brak danych rejestracyjnych HTX.")
            return

        self._add_paragraph("")
        self._add_heading("Profil użytkownika HTX", level=3)

        # Dla każdego UID wygeneruj kartę profilu
        for uid in sorted(htx_data.keys()):
            user_data = htx_data[uid]
            if uid.startswith("row_"):
                # Brak UID — pokaż jako anonimowy wiersz
                self._add_paragraph("")
                self._add_heading(f"Wiersz bez UID ({uid})", level=4)
                rows = [[k, v, "register_1"] for k, v in user_data.items()]
                self._add_table(["Pole", "Wartość", "Źródło"], rows, max_rows=50)
                continue

            self._render_htx_user_card(r, uid, user_data)

        # Portfele i transakcje wyświetlane są per użytkownik w _render_htx_user_card

    def _render_htx_user_card(self, r: ExtractedIdentifiers, uid: str, user_data: Dict[str, str]):
        """Renderuje kartę pojedynczego użytkownika HTX: dane → zdjęcia → portfele."""
        self._add_heading(f"Użytkownik HTX — UID: {uid}", level=4)
        print(f"    [DOCX] UID {uid}: dane={len(user_data)} pól, zdjęcia={len(r.certified_photos.get(uid,[]))}, portfele={len(r.wallet_addresses_by_user.get(uid,[]))}")

        # --- 1. TABELA DANYCH OSOBOWYCH ---
        field_labels = {
            "uid": "UID użytkownika",
            "name": "Imię i nazwisko",
            "email": "E-mail",
            "phone": "Telefon",
            "idcard": "Numer dokumentu",
            "gmt_created": "Data utworzenia konta",
            "country": "Kraj rejestracji",
            "balance": "Saldo",
            "bankcard": "Karta bankowa",
            "alipay": "Alipay",
            "wechat": "WeChat",
        }

        priority_fields = ["uid", "gmt_created", "name", "email", "phone", "idcard", "country"]
        profile_rows = []

        def _fmt_val(val):
            s = str(val) if val is not None else "(puste)"
            if s.endswith(".0"):
                s = s[:-2]
            return s

        for field_key in priority_fields:
            if field_key in user_data:
                label = field_labels.get(field_key, field_key)
                profile_rows.append([label, _fmt_val(user_data[field_key]), "register_1"])

        for field_key, val in user_data.items():
            if field_key in priority_fields or field_key == "user_address" or field_key == "balances_list":
                continue
            label = field_labels.get(field_key, field_key)
            profile_rows.append([label, _fmt_val(val), "register_1"])

        # Salda szczegółowe z balance_1 — w jednej komórce, po enterach
        if "balances_list" in user_data:
            bal_items = user_data["balances_list"].split(" | ")
            bal_multiline = "\n".join(bal_items)
            profile_rows.append(["Salda szczegółowe", bal_multiline, "balance_1"])

        if profile_rows:
            self._add_table(["Pole", "Wartość", "Źródło"], profile_rows, max_rows=50)

        # --- 2. LOGOWANIA ---
        self._render_htx_logins_for_user(r, uid)

        # --- 3. ZDJĘCIA CERTYFIKOWANE ---
        self._render_htx_photos(r, uid)

        # --- 3. TABELA PORTFELI (osobna, na końcu) ---
        self._add_paragraph("")
        wallets = r.wallet_addresses_by_user.get(uid, [])
        if wallets:
            self._add_paragraph("Adresy portfeli kryptowalutowych:", bold=True)
            wallet_rows = [[str(i+1), addr, "register_1"] for i, addr in enumerate(sorted(wallets))]
            self._add_table(["Lp.", "Adres portfela", "Źródło"], wallet_rows, max_rows=20)
        else:
            self._add_paragraph("Brak przypisanych adresów portfeli.", color=RGBColor(0x80, 0x80, 0x80))

        self._add_paragraph("")  # odstęp między użytkownikami

    def _render_htx_logins_for_user(self, r: ExtractedIdentifiers, uid: str):
        """Renderuje logowania HTX dla konkretnego UID.

        Pokazuje:
        - Tabelkę logowań (max 50 wierszy): Czas, Terminal, IP, Źródło
        - Unikalne IP z zakresem czasowym (min/max daty) per IP
        """
        records = r.login_records.get(uid, [])
        if not records:
            return

        self._add_paragraph("")
        self._add_heading(f"Historia logowania", level=4)

        # --- Tabela szczegółowa logowań (max 50) ---
        login_rows = []
        for rec in records[:50]:
            login_rows.append([
                rec.get("time", ""),
                rec.get("terminal", ""),
                rec.get("ip", ""),
                "login_1",
            ])
        self._add_table(["Czas logowania", "Terminal", "Adres IP", "Źródło"], login_rows, max_rows=50)
        if len(records) > 50:
            self._add_paragraph(
                f"... i {len(records) - 50} więcej logowań (pełna lista w JSON).",
                color=RGBColor(0x80, 0x80, 0x80))

        # --- Tabela unikalnych IP z zakresem czasowym ---
        from collections import defaultdict
        ip_records = defaultdict(list)
        for rec in records:
            ip = rec.get("ip", "")
            if ip:
                ip_records[ip].append(rec.get("time", ""))

        if ip_records:
            self._add_paragraph("Unikalne adresy IP z zakresem czasowym:", bold=True)
            ip_rows = []
            for ip, times in sorted(ip_records.items()):
                valid_times = [t for t in times if t]
                if valid_times:
                    time_from = min(valid_times)
                    time_to = max(valid_times)
                    ip_rows.append([ip, str(len(times)), time_from, time_to, "login_1"])
                else:
                    ip_rows.append([ip, str(len(times)), "(brak daty)", "(brak daty)", "login_1"])
            self._add_table(
                ["Adres IP", "Liczba logowań", "Pierwsze logowanie", "Ostatnie logowanie", "Źródło"],
                ip_rows, max_rows=50)

        self._add_paragraph("")

    def _render_htx_photos(self, r: ExtractedIdentifiers, uid: str):
        """Renderuje zdjęcia certyfikowane dla danego UID w tabeli.

        Jeśli zdjęcie >1MB, nie osadza się w DOCX (zbyt duże), tylko pokazuje ścieżkę.
        """
        photos = r.certified_photos.get(uid, [])
        if not photos:
            return

        # Sprawdź rozmiary — jeśli suma >5MB, pokaż tylko listę
        total_size = sum(os.path.getsize(p) for p in photos if os.path.exists(p))
        if total_size > 5 * 1024 * 1024:
            self._add_paragraph("")
            self._add_paragraph(f"Zdjęcia certyfikowane ({len(photos)} plików, {total_size/1024/1024:.1f} MB — za duże do osadzenia w raporcie):", bold=True)
            rows = [[str(i+1), os.path.basename(p), p, "certified_photos"] for i, p in enumerate(photos)]
            self._add_table(["Lp.", "Nazwa pliku", "Ścieżka", "Źródło"], rows, max_rows=20)
            return

        self._add_paragraph(f"Zdjęcia certyfikowane ({len(photos)}):", bold=True)
        # Tabela z zdjęciami — max 3 kolumny
        n_cols = min(3, len(photos))
        n_rows = (len(photos) + n_cols - 1) // n_cols
        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.allow_autofit = False

        col_w = Inches(PAGE_WIDTH_INCHES / n_cols)
        for i in range(n_cols):
            table.columns[i].width = col_w

        for img_idx, img_path in enumerate(photos):
            row_idx = img_idx // n_cols
            col_idx = img_idx % n_cols
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            if os.path.exists(img_path):
                # Jeśli plik >1MB, nie osadzaj — za duży
                file_size = os.path.getsize(img_path)
                if file_size > 1024 * 1024:
                    cell.text = f"{os.path.basename(img_path)}\n({file_size/1024/1024:.1f} MB — za duże)"
                else:
                    try:
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        img_width = Inches(1.2) if n_cols >= 3 else Inches(1.8)
                        run.add_picture(img_path, width=img_width)
                    except Exception as e:
                        cell.text = f"[!] Błąd: {e}"
            else:
                cell.text = "Brak pliku"

    def _render_htx_wallets(self, r: ExtractedIdentifiers):
        """Globalna sekcja portfeli HTX (podsumowanie wszystkich UID)."""
        if not r.wallet_addresses_by_user:
            return

        self._add_heading("Portfele kryptowalutowe (podsumowanie)", level=3)
        rows = []
        for uid, addrs in sorted(r.wallet_addresses_by_user.items()):
            rows.append([str(uid), ", ".join(sorted(addrs)), str(len(addrs)), "register_1"])
        self._add_table(["UID", "Adresy portfeli", "Liczba", "Źródło"], rows, max_rows=50)

    # ========================================================================
    # KONIEC HTX
    # ========================================================================

    def _render_customer_info(self, r: ExtractedIdentifiers):
        if not r.customer_info_sections:
            return
        self._add_heading("Customer Information (Informacje o użytkowniku)", level=3)
        for sec_name, data in r.customer_info_sections.items():
            if not data:
                continue
            display_name = SECTION_TRANSLATIONS.get(sec_name, sec_name)
            self._add_heading(display_name, level=4)
            if sec_name == "Basic Information":
                rows = self._sort_basic_info(data)
            elif sec_name == "HTX Register":
                # HTX Register ma strukturę: {uid: {col: val, ...}}
                # Dla Binance (exchange != htx) renderujemy stare stylem
                rows = []
                for uid_key, row_data in data.items():
                    for col, val in row_data.items():
                        if val is not None and str(val).strip() != "":
                            rows.append([str(col), str(val)])
                        else:
                            rows.append([str(col), "(puste)"])
            else:
                rows = []
                for col, val in data.items():
                    if val is not None and str(val).strip() != "":
                        rows.append([str(col), str(val)])
                    else:
                        rows.append([str(col), "(puste)"])
            self._add_table(["Pole", "Wartość"], rows, max_rows=100)

    def _render_kyc(self, r: ExtractedIdentifiers):
        self._add_heading("KYC Documents (Dokumenty KYC)", level=3)
        if r.kyc_images:
            n_cols = 2
            n_rows = (len(r.kyc_images) + n_cols - 1) // n_cols
            table = self.doc.add_table(rows=n_rows, cols=n_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            table.allow_autofit = False
            col_w = Inches(PAGE_WIDTH_INCHES / n_cols)
            for i in range(n_cols):
                table.columns[i].width = col_w
            for img_idx, img_path in enumerate(r.kyc_images):
                row_idx = img_idx // n_cols
                col_idx = img_idx % n_cols
                cell = table.cell(row_idx, col_idx)
                cell.text = ""
                if os.path.exists(img_path):
                    try:
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(img_path, width=Inches(2.8))
                    except Exception as e:
                        cell.text = f"[!] Błąd: {e}"
                else:
                    cell.text = "Brak pliku"
        else:
            self._add_paragraph("Brak osadzonych obrazków w arkuszu KYC Documents.")

    def _render_assets_overview(self, r: ExtractedIdentifiers):
        self._add_heading("Assets Overview (Przegląd aktywów)", level=3)
        if r.estimate_total_btc:
            self._add_paragraph(f"Estimate Total Balance (BTC): {r.estimate_total_btc}", bold=True)
        if r.estimate_total_usdt:
            self._add_paragraph(f"Estimate Total Balance (USD): ≈${r.estimate_total_usdt}", bold=True)
        if r.asset_balances:
            bal_rows = []
            for b in r.asset_balances:
                waluta = b.currency_code
                if b.currency_name and b.currency_name != "—":
                    waluta = f"{b.currency_code} ({b.currency_name})"
                bal_rows.append([
                    waluta,
                    b.all_positions,
                    b.btc_equivalent,
                    b.wallet_type,
                ])
            self._add_table(
                ["Waluta", "Saldo", "Wartość BTC", "Portfel"],
                bal_rows, max_rows=100)
        else:
            self._add_paragraph("Brak danych o saldach walut (wszystkie salda to 0 lub arkusz nie zawiera tabeli).")

    def _render_currency_flows(self, r: ExtractedIdentifiers):
        """Jedna sekcja per waluta — wszystkie transakcje ze wszystkich arkuszy w jednym miejscu.

        Uwaga: Transakcje z arkuszy Fiat (Fiat Deposit, Fiat Trades) są wyświetlane w tabeli
        dla czytelności (widać źródło zakupu), ale NIE są liczone do bilansu, ponieważ
        ten sam ruch jest już odzwierciedlony w Spot Asset Log / Funding Asset Log.
        """
        all_txns = (r.spot_transactions + r.funding_transactions +
                    r.deposit_transactions + r.withdrawal_transactions +
                    r.fiat_deposit_transactions + r.fiat_trade_transactions)
        if not all_txns and not r.asset_balances:
            return

        self._add_heading("Wykaz ruchów walut (wszystkie źródła)", level=3)
        self._add_paragraph(
            "Poniższa analiza łączy dane ze wszystkich arkuszy transakcyjnych: Spot Asset Log, Funding Asset Log, "
            "Deposit History, Withdrawal History, Fiat Deposit oraz Fiat Trades. Dla każdej waluty pokazano pełen bilans ruchów. "
            "Transakcje oznaczone [FIAT] w kolumnie Powód pochodzą z arkuszy fiat. "
            "Wartości w nawiasie (*) to duplikaty — ten sam ruch jest już uwzględniony w logach Spot/Funding i NIE wpływa na bilans. "
            "Wartości bez nawiasu to unikalne transakcje fiat (brak duplikatu w innych logach) i SĄ wliczane do bilansu.",
            color=RGBColor(0x60, 0x60, 0x60))

        # Zbierz wszystkie waluty: z transakcji + z Assets Overview
        currencies_from_txns = set()
        currencies_from_balances = set()

        confirmed_all = [t for t in all_txns if not _is_pending_transaction(t)]
        pending_all = [t for t in all_txns if _is_pending_transaction(t)]

        # === BILANS: wszystkie transakcje, ale fiat-duplikaty pomijamy ===
        confirmed_for_balance = [t for t in confirmed_all if not _is_fiat_duplicate(t, confirmed_all)]

        # Grupowanie transakcji per waluta (do wyświetlenia — wszystkie, w tym fiat)
        txns_by_currency: Dict[str, List[AssetTransaction]] = {}
        for t in confirmed_all:
            curr = t.currency.upper() if t.currency else "UNKNOWN"
            currencies_from_txns.add(curr)
            if curr not in txns_by_currency:
                txns_by_currency[curr] = []
            txns_by_currency[curr].append(t)

        # Waluty z Assets Overview
        balance_map: Dict[str, AssetBalance] = {}
        for b in r.asset_balances:
            curr = b.currency_code.upper()
            currencies_from_balances.add(curr)
            balance_map[curr] = b

        all_currencies = sorted(currencies_from_txns | currencies_from_balances)

        if pending_all:
            self._add_paragraph(
                f"Pominięto {len(pending_all)} transakcji pending/niepotwierdzonych (nieuwzględnionych w bilansie).",
                color=RGBColor(0x80, 0x60, 0x00))

        if not all_currencies:
            self._add_paragraph("Brak danych o ruchach walut.")
            return

        # Najpierw skrótowa tabela wszystkich walut (bilans zgodny ze szczegółami — bez duplikatów fiat)
        summary_rows = []
        for curr in all_currencies:
            txns_bal = [t for t in txns_by_currency.get(curr, []) if not _is_fiat_duplicate(t, confirmed_all)]
            total_in = sum(_to_float(t.change) for t in txns_bal if _to_float(t.change) > 0)
            total_out = sum(abs(_to_float(t.change)) for t in txns_bal if _to_float(t.change) < 0)
            netto = total_in - total_out
            bal = balance_map.get(curr)
            bal_str = bal.all_positions if bal else "(brak)"
            diff_str = ""
            if bal:
                bal_f = _to_float(bal.all_positions)
                diff = bal_f - netto
                if abs(diff) > 1e-8:
                    diff_str = _fmt_signed(diff)
            summary_rows.append([
                curr,
                f"+{_fmt(total_in)}",
                f"-{_fmt(total_out)}",
                _fmt_signed(netto),
                bal_str,
                diff_str,
            ])

        self._add_table(
            ["Waluta", "Przychody (wszystkie)", "Rozchody (wszystkie)", "Netto", "Saldo z Assets Overview", "Różnica"],
            summary_rows)

        self._add_paragraph(
            "ℹ️ Kolumna 'Różnica' pokazuje różnicę między bilansem przepływów (bez duplikatów fiat) a saldem z Assets Overview. "
            "Jeżeli różnica jest niezerowa, oznacza to że część środków została przeniesiona między portfelami "
            "lub znajduje się w innych produktach (Futures, Earn, Margin, Pool).",
            color=RGBColor(0x00, 0x60, 0x80))
        self._add_paragraph("")

        # Szczegóły per waluta
        for curr in all_currencies:
            txns = txns_by_currency.get(curr, [])
            bal = balance_map.get(curr)

            # Sprawdź czy w tej walucie są transakcje fiat z obcym User ID
            # Właściciel to TYLKO ID z Customer Information (Basic Information)
            # Każde inne ID w transakcjach fiat = potencjalne zasilenie z zewnątrz
            foreign_in_curr = set()
            for t in txns:
                if t.wallet_type == "Fiat" and t.user_id and t.user_id not in r.user_ids:
                    foreign_in_curr.add(t.user_id)

            # Podsumowanie per waluta (bez duplikatów fiat)
            txns_bal = [t for t in txns if not _is_fiat_duplicate(t, confirmed_all)]
            total_in = sum(_to_float(t.change) for t in txns_bal if _to_float(t.change) > 0)
            total_out = sum(abs(_to_float(t.change)) for t in txns_bal if _to_float(t.change) < 0)
            netto = total_in - total_out

            times = [t.time for t in txns if t.time]
            time_range_str = ""
            if times:
                time_range_str = f" (zakres: {min(times)[:19]} → {max(times)[:19]})"

            self._add_paragraph(f"Waluta {curr}:{time_range_str}", bold=True)

            if foreign_in_curr:
                self._add_paragraph(
                    f" ⚠️ UWAGA: Transakcje fiat w tej walucie pochodzą od innego User ID: {', '.join(sorted(foreign_in_curr))}. "
                    f"Może to oznaczać zasilenie konta z zewnętrznego źródła (inny użytkownik / konto powiązane).",
                    color=RGBColor(0xC0, 0x00, 0x00))

            # Zlicz FAIL w tej walucie (z fiat)
            fail_in_curr = sum(1 for t in txns if t.wallet_type == "Fiat" and t.reason and "Status: fail" in t.reason.lower())
            if fail_in_curr > 0:
                self._add_paragraph(
                    f" ℹ️ W tej walucie wykryto {fail_in_curr} nieudanych transakcji fiat (FAIL) — pominięte w bilansie.",
                    color=RGBColor(0x80, 0x60, 0x00))

            if txns_bal:
                summary_row = [[
                    f"+{_fmt(total_in)}",
                    f"-{_fmt(total_out)}",
                    _fmt_signed(netto),
                    str(sum(1 for t in txns_bal if _to_float(t.change) > 0)),
                    str(sum(1 for t in txns_bal if _to_float(t.change) < 0)),
                ]]
                self._add_table(
                    ["Przychody", "Rozchody", "Saldo", "L. przych.", "L. rozch."],
                    summary_row)

                if bal:
                    bal_f = _to_float(bal.all_positions)
                    if abs(netto - bal_f) > 1e-8:
                        diff = bal_f - netto
                        self._add_paragraph(
                            f" ℹ️ Różnica między logiem ({_fmt_signed(netto)}) a Assets Overview ({_fmt_signed(bal_f)}): {_fmt_signed(diff)} {curr}. "
                            f"Brakujące transakcje (depozyty, wypłaty, transfery między portfelami) znajdują się w innych arkuszach lub produktach.",
                            color=RGBColor(0x00, 0x60, 0x80))

                if netto < -1e-12:
                    self._add_paragraph(
                        f" ⚠️ UWAGA: Ujemne saldo w logu ({_fmt_signed(netto)}). "
                        f"Oznacza to że w okresie objętym raportem rozchody przewyższyły przychody. "
                        f"Nie oznacza to debetu — brakuje tu depozytów/wypłat z innych źródeł.",
                        color=RGBColor(0xC0, 0x00, 0x00))

                # Tabela transakcji — WSZYSTKIE (w tym fiat, oznaczone jako duplikat lub unikalne)
                txn_rows = []
                for t in sorted(txns, key=lambda x: x.time or ""):
                    chg = _to_float(t.change)
                    if abs(chg) < 1e-12 and t.wallet_type != "Fiat":
                        continue

                    is_dup = _is_fiat_duplicate(t, confirmed_all)
                    chg_str = _fmt_signed(chg)
                    # Dla fiat-duplikatów: nawias + gwiazdka
                    if t.wallet_type == "Fiat" and is_dup:
                        chg_str = f"({_fmt_signed(chg)})*"
                    # Dla fiat-unikalnych: bez nawiasu (liczy się do bilansu)
                    elif t.wallet_type == "Fiat" and not is_dup:
                        chg_str = _fmt_signed(chg)

                    reason_display = t.reason if t.reason else "—"
                    source_display = t.source_sheet if t.source_sheet else t.wallet_type
                    txn_rows.append([
                        t.time[:19] if t.time else "",
                        chg_str,
                        reason_display,
                        source_display,
                        t.transaction_id[:20] if t.transaction_id else "—",
                    ])

                if txn_rows:
                    self._add_table(
                        ["Czas", "Zmiana", "Powód", "Źródło", "TxID"],
                        txn_rows,
                        col_widths=[Inches(1.0), Inches(0.8), Inches(2.8), Inches(1.1), Inches(1.3)])

                    # Dodaj legendę, jeśli były duplikaty fiat
                    has_fiat_dup = any(t.wallet_type == "Fiat" and _is_fiat_duplicate(t, confirmed_all) for t in txns)
                    has_fiat_unique = any(t.wallet_type == "Fiat" and not _is_fiat_duplicate(t, confirmed_all) for t in txns)
                    if has_fiat_dup:
                        self._add_paragraph(
                            "* Wartość w nawiasie to duplikat fiat — pokazana informacyjnie, nie wliczana do bilansu (ten sam ruch jest w logach Spot/Funding).",
                            color=RGBColor(0x80, 0x80, 0x80))
                    if has_fiat_unique:
                        self._add_paragraph(
                            "Transakcja unikalna z arkusza Fiat (brak duplikatu w innych logach) — WYLICZANA do bilansu.",
                            color=RGBColor(0x00, 0x60, 0x80))
            else:
                # Waluta w Assets Overview ale bez transakcji (lub tylko fiat)
                bal_str = bal.all_positions if bal else "(brak)"
                wallet_str = f" ({bal.wallet_type})" if bal and bal.wallet_type else ""
                self._add_paragraph(
                    f"Brak transakcji we wszystkich logach. Saldo z Assets Overview: {_fmt_signed(_to_float(bal_str))} {curr}{wallet_str}.",
                    color=RGBColor(0x60, 0x60, 0x60))

            self._add_paragraph("")

    def generate(self, reports: List[ExtractedIdentifiers], file_map: Dict[str, str]):
        # ===== STRONA TYTUŁOWA =====
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("RAPORT ANALIZY")
        run.bold = True
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run("Raportów giełdowych kryptowalutowych")
        run2.font.size = Pt(18)

        self.doc.add_paragraph()
        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.add_run(f"Data wygenerowania: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        self.doc.add_page_break()

        # ===== SPIS TREŚCI =====
        self._add_heading("Spis treści", level=1)
        toc_items = ["1. Podsumowanie", "2. Szczegółowa analiza poszczególnych plików"]
        for i, r in enumerate(reports, 1):
            toc_items.append(f"  2.{i}. {r.source_file}")
        if len(reports) >= 2:
            toc_items.append("3. Porównanie między raportami")
        toc_items.append("4. Pełna lista identyfikatorów")
        for item in toc_items:
            self.doc.add_paragraph(item, style='List Bullet')
        self.doc.add_page_break()

        # ===== 1. PODSUMOWANIE =====
        self._add_heading("1. Podsumowanie", level=1)
        self._add_paragraph(f"Liczba przeanalizowanych plików: {len(reports)}")
        self._add_paragraph(f"Liczba giełd: {len(set(file_map.values()))}")

        summary_rows = []
        for r in reports:
            all_from = []
            all_to = []
            for tr in r.time_ranges.values():
                all_from.append(tr.get("from", ""))
                all_to.append(tr.get("to", ""))
            time_summary = ""
            if all_from and all_to:
                time_summary = f"{min(all_from)} -> {max(all_to)}"

            # Dla każdego UID osobny wiersz (HTX: wiele użytkowników per plik)
            uids = sorted(r.user_ids) if r.user_ids else [""]
            for uid in uids:
                wallet_count = len(r.wallet_addresses_by_user.get(uid, [])) if uid else len(r.wallet_addresses)

                # Dla HTX: wyciągnij dane per UID z customer_info_sections
                email_per_uid = ""
                phone_per_uid = ""
                name_per_uid = ""
                balance_per_uid = ""
                if r.exchange == "htx" and uid and "HTX Register" in r.customer_info_sections:
                    user_data = r.customer_info_sections["HTX Register"].get(uid, {})
                    email_per_uid = user_data.get("email", "")
                    phone_per_uid = user_data.get("phone", "")
                    name_per_uid = user_data.get("name", "")
                    balance_per_uid = user_data.get("balance", "")

                summary_rows.append([
                    r.source_file,
                    r.exchange.upper(),
                    str(len(r.parsed_sheets)),
                    uid if uid else "(brak)",
                    name_per_uid,
                    email_per_uid,
                    phone_per_uid,
                    str(wallet_count),
                    balance_per_uid,
                    str(len(r.asset_balances)),
                    time_summary,
                ])
        self._add_table(
            ["Plik", "Giełda", "Arkusze", "UID", "Imię i nazwisko", "E-mail", "Telefon", "Portfele", "Saldo", "Salda walut", "Zakres czasowy"],
            summary_rows)
        self.doc.add_page_break()

        # ===== 2. SZCZEGÓŁY KAŻDEGO PLIKU =====
        self._add_heading("2. Szczegółowa analiza poszczególnych plików", level=1)

        for idx, r in enumerate(reports, 1):
            self._add_heading(f"2.{idx}. {r.source_file}", level=2)
            self._add_paragraph(f"Giełda: {r.exchange.upper()}")
            if r.user_ids:
                uid_str = ", ".join(sorted(r.user_ids))
                self._add_paragraph(f"ID użytkownika (właściciel konta): {uid_str}", bold=True, color=RGBColor(0x00, 0x00, 0x80))
            self._add_paragraph(
                f"Przeanalizowane arkusze ({len(r.parsed_sheets)}): {', '.join(r.parsed_sheets)}")
            if r.unknown_sheets:
                self._add_paragraph(
                    f"Nieznane arkusze: {', '.join(r.unknown_sheets)}",
                    color=RGBColor(0xC0, 0x00, 0x00))

            all_from = []
            all_to = []
            for tr in r.time_ranges.values():
                all_from.append(tr.get("from", ""))
                all_to.append(tr.get("to", ""))
            if all_from and all_to:
                self._add_paragraph(
                    f"Globalny zakres czasowy konta: {min(all_from)} → {max(all_to)}",
                    bold=True, color=RGBColor(0x00, 0x40, 0x80))
            self._add_paragraph("")

            # ===== HTX: PROFIL UŻYTKOWNIKA =====
            if r.exchange == "htx":
                self._render_htx_profile(r)
            else:
                # Binance: klasyczna struktura
                for sheet_name in r.parsed_sheets:
                    if sheet_name == "Customer Information":
                        self._render_customer_info(r)
                    elif sheet_name == "KYC Documents":
                        self._render_kyc(r)
                    elif sheet_name == "Assets Overview":
                        self._render_assets_overview(r)
                    elif sheet_name == "register_1":
                        self._render_customer_info(r)

                # Nowa sekcja transakcyjna — per waluta, wszystkie źródła razem
                self._render_currency_flows(r)

            self.doc.add_page_break()

        # ===== 3. PORÓWNANIE =====
        if len(reports) >= 2:
            self._add_heading("3. Porównanie między raportami", level=1)
            comp = ReportComparator(reports)
            result = comp.compare()
            common = result.get("common", {})
            if common:
                self._add_heading("Wspólne identyfikatory (potencjalne powiązania):", level=2)
                self._add_paragraph(
                    "Poniższe identyfikatory występują w co najmniej dwóch raportach. "
                    "Mogą wskazywać na powiązanie między kontami.",
                    color=RGBColor(0xC0, 0x00, 0x00))
                for field_name, entries in common.items():
                    self._add_heading(f"[{field_name}] — {len(entries)} wspólnych", level=3)
                    rows = []
                    for i, e in enumerate(entries):
                        files_str = ", ".join(e["files"])
                        time_str = ""
                        if "time_context" in e:
                            parts = []
                            for tc in e["time_context"]:
                                if tc["ranges"]:
                                    parts.append(f"{tc['file']}: {tc['ranges'][0]}")
                            time_str = "; ".join(parts)
                        rows.append([str(i+1), str(e["value"]), files_str, time_str])
                    self._add_table(["Lp.", "Wartość", "Pliki", "Zakres czasowy"], rows, max_rows=30)
            else:
                self._add_paragraph("Nie znaleziono wspólnych identyfikatorów między raportami.")

            # --- WSPÓLNE IP MIĘDZY KONTAMI HTX ---
            self._render_shared_ips_comparison(reports)

            self.doc.add_page_break()

        # ===== 4. PEŁNA LISTA IDENTYFIKATORÓW =====
        self._add_heading("4. Pełna lista identyfikatorów", level=1)
        self._add_paragraph("Szczegółowe dane w formacie JSON zostały zapisane w pliku parsed_report.json.")

        for r in reports:
            self._add_heading(f"{r.source_file}", level=2)

            # Dla HTX: sekcja 4 zawiera tylko dane NIEprezentowane w sekcji 2.4
            if r.exchange == "htx":
                # Tylko zdjęcia certyfikowane (bo są tylko tu, per UID ze ścieżkami)
                if r.certified_photos:
                    self._add_paragraph(f"Zdjęcia certyfikowane ({sum(len(v) for v in r.certified_photos.values())} zdjęć):", bold=True)
                    photo_rows = []
                    for uid_val, paths in sorted(r.certified_photos.items()):
                        for p in paths:
                            photo_rows.append([str(uid_val), os.path.basename(p), p])
                    self._add_table(["UID", "Nazwa pliku", "Ścieżka"], photo_rows, max_rows=50)
                # Pozostałe dane (imiona, emaile, portfele) są już w sekcji 2.4 — pomijamy
                continue

            # Dla Binance: pełna lista identyfikatorów (jak poprzednio)
            if r.user_ids:
                self._add_paragraph("")
                self._add_paragraph(f"ID właściciela konta ({len(r.user_ids)}):", bold=True)
                rows = [[str(i+1), str(item)] for i, item in enumerate(sorted(r.user_ids))]
                self._add_table(["Lp.", "Wartość"], rows, max_rows=20)
            if r.related_user_ids:
                self._add_paragraph("")
                self._add_paragraph(f"ID powiązanych użytkowników ({len(r.related_user_ids)}):", bold=True)
                rows = [[str(i+1), str(item)] for i, item in enumerate(sorted(r.related_user_ids))]
                self._add_table(["Lp.", "Wartość"], rows, max_rows=20)

            id_sections = [
                ("E-maile", r.emails), ("Numery telefonów", r.phones),
                ("Adresy IP", r.ips), ("Adresy portfeli (krypto)", r.wallet_addresses),
                ("TXID (hash transakcji)", r.txids), ("BIN karty", r.card_bins),
                ("Ostatnie 4 cyfry karty", r.card_last4),
                ("Karty płatnicze (sformatowane)", r.formatted_cards),
                ("IBAN", r.ibans),
                ("Numery kont", r.account_numbers), ("ID urządzeń", r.device_ids),
                ("ID Fvideo", r.fvideo_ids), ("UUID BNC", r.bnc_uuids),
                ("ID zamówień", r.order_ids), ("ID kontrahentów", r.counterparty_ids),
                ("ID transakcji", r.transaction_ids), ("Imiona i nazwiska", r.names),
                ("Narodowości", r.nationalities), ("Numery dokumentów tożsamości", r.id_numbers),
                ("Lokalizacje GEO", r.geolocations), ("Przeglądarki / User Agent", r.browsers),
            ]
            for title, items in id_sections:
                if items:
                    self._add_paragraph("")
                    self._add_paragraph(f"{title} ({len(items)}):", bold=True)
                    rows = [[str(i+1), str(item)] for i, item in enumerate(sorted(items))]
                    self._add_table(["Lp.", "Wartość"], rows, max_rows=20)

        self.doc.save(self.output_path)

    def _render_shared_ips_comparison(self, reports: List[ExtractedIdentifiers]):
        """Wykrywa i wyświetla adresy IP współdzielone między kontami HTX.

        Jeśli ten sam IP występuje w logowaniach więcej niż jednego użytkownika
        (w tym samym lub różnym raporcie), generuje osobną tabelkę ostrzegawczą.
        """
        # Zbierz wszystkie IP per (raport, uid)
        ip_to_users: Dict[str, List[Dict]] = {}
        for r in reports:
            if r.exchange != "htx":
                continue
            for uid, records in r.login_records.items():
                for rec in records:
                    ip = rec.get("ip", "").strip()
                    if not ip or not is_valid_ip(ip):
                        continue
                    if ip not in ip_to_users:
                        ip_to_users[ip] = []
                    # Sprawdź czy ten (raport, uid) już jest
                    existing = [e for e in ip_to_users[ip] if e["file"] == r.source_file and e["uid"] == uid]
                    if not existing:
                        ip_to_users[ip].append({
                            "file": r.source_file,
                            "uid": uid,
                            "times": [rec.get("time", "")],
                        })
                    else:
                        existing[0]["times"].append(rec.get("time", ""))

        # Filtruj tylko IP używane przez >1 użytkownika (w tym samym lub innym raporcie)
        shared_ips = {ip: entries for ip, entries in ip_to_users.items() if len(entries) > 1}
        if not shared_ips:
            return

        self._add_heading("Współdzielone adresy IP (HTX)", level=3)
        self._add_paragraph(
            "Poniższe adresy IP występują w logowaniach więcej niż jednego użytkownika. "
            "Może to oznaczać współdzielone urządzenie, sieć VPN/proxy, lub powiązanie między kontami.",
            color=RGBColor(0xC0, 0x00, 0x00))

        rows = []
        for ip, entries in sorted(shared_ips.items()):
            # Zbuduj listę użytkowników per plik
            users_per_file: Dict[str, List[str]] = {}
            for e in entries:
                if e["file"] not in users_per_file:
                    users_per_file[e["file"]] = []
                users_per_file[e["file"]].append(e["uid"])

            user_strs = []
            for fname, uids in users_per_file.items():
                user_strs.append(f"{fname}: {', '.join(sorted(set(uids)))}")

            # Zakres czasowy dla tego IP
            all_times = []
            for e in entries:
                all_times.extend([t for t in e["times"] if t])
            time_range = ""
            if all_times:
                time_range = f"{min(all_times)} → {max(all_times)}"

            rows.append([ip, "; ".join(user_strs), time_range])

        self._add_table(
            ["Adres IP", "Użytkownicy (plik: UID)", "Zakres czasowy"],
            rows, max_rows=50)

    # ========================================================================
    # HTX: NOWE SEKCJE — Logowania, Urządzenia, Transakcje
    # ========================================================================

    def _render_htx_logins(self, r: ExtractedIdentifiers):
        """Renderuje logowania HTX w DOCX — max 50 wierszy."""
        login_sheets = [s for s in r.parsed_sheets if s == "login_1"]
        if not login_sheets and not r.ips and not r.geolocations and not r.browsers:
            return
        self._add_heading("Historia logowania", level=3)
        # Zakres czasowy
        for sheet_name in login_sheets:
            if sheet_name in r.time_ranges:
                tr = r.time_ranges[sheet_name]
                self._add_paragraph(f"Zakres czasowy: {tr['from']} → {tr['to']}")
        # IP
        if r.ips:
            self._add_paragraph("")
            self._add_paragraph(f"Unikalne adresy IP ({len(r.ips)}):", bold=True)
            rows = [[str(i+1), ip, "login_1"] for i, ip in enumerate(sorted(r.ips)[:50])]
            self._add_table(["Lp.", "Adres IP", "Źródło"], rows, max_rows=50)
            if len(r.ips) > 50:
                self._add_paragraph(f"... i {len(r.ips)-50} więcej (pełna lista w JSON).", color=RGBColor(0x80, 0x80, 0x80))
        # Geolokalizacje
        if r.geolocations:
            self._add_paragraph("")
            self._add_paragraph(f"Lokalizacje ({len(r.geolocations)}):", bold=True)
            rows = [[str(i+1), loc, "login_1"] for i, loc in enumerate(sorted(r.geolocations)[:50])]
            self._add_table(["Lp.", "Lokalizacja", "Źródło"], rows, max_rows=50)
            if len(r.geolocations) > 50:
                self._add_paragraph(f"... i {len(r.geolocations)-50} więcej (pełna lista w JSON).", color=RGBColor(0x80, 0x80, 0x80))
        # Przeglądarki
        if r.browsers:
            self._add_paragraph("")
            self._add_paragraph(f"Przeglądarki / User Agent ({len(r.browsers)}):", bold=True)
            rows = [[str(i+1), br, "login_1"] for i, br in enumerate(sorted(r.browsers)[:50])]
            self._add_table(["Lp.", "Przeglądarka", "Źródło"], rows, max_rows=50)
            if len(r.browsers) > 50:
                self._add_paragraph(f"... i {len(r.browsers)-50} więcej (pełna lista w JSON).", color=RGBColor(0x80, 0x80, 0x80))

    def _render_htx_devices(self, r: ExtractedIdentifiers):
        """Renderuje skrót urządzeń HTX (DeviceFP_1)."""
        device_sheets = [s for s in r.parsed_sheets if s == "DeviceFP_1"]
        if not device_sheets and not r.device_ids:
            return
        self._add_heading("Urządzenia zatwierdzone (skrót)", level=3)
        if r.device_ids:
            self._add_paragraph(f"Liczba unikalnych urządzeń / odcisków: {len(r.device_ids)}")
            rows = [[str(i+1), dev[:80], "DeviceFP_1"] for i, dev in enumerate(sorted(r.device_ids)[:20])]
            if rows:
                self._add_table(["Lp.", "ID urządzenia", "Źródło"], rows, max_rows=20)
            if len(r.device_ids) > 20:
                self._add_paragraph(f"... i {len(r.device_ids)-20} więcej (pełna lista w JSON).", color=RGBColor(0x80, 0x80, 0x80))

    def _render_htx_transactions(self, r: ExtractedIdentifiers):
        """Renderuje transakcje HTX w DOCX — max 50 wierszy per typ.
        Pełne dane w pliku JSON.
        """
        has_txns = (r.deposit_transactions or r.withdrawal_transactions or
                    r.spot_transactions or r.funding_transactions)
        if not has_txns:
            return
        self._add_heading("Transakcje HTX", level=3)

        # Liczby
        counts = []
        if r.deposit_transactions:
            counts.append(f"Depozyty: {len(r.deposit_transactions)}")
        if r.withdrawal_transactions:
            counts.append(f"Wypłaty: {len(r.withdrawal_transactions)}")
        if r.spot_transactions:
            counts.append(f"Trade/Spot: {len(r.spot_transactions)}")
        if r.funding_transactions:
            counts.append(f"Funding: {len(r.funding_transactions)}")
        self._add_paragraph(" | ".join(counts))

        # Deposit — max 50
        if r.deposit_transactions:
            self._add_paragraph("")
            self._add_heading("Depozyty (pierwsze 50)", level=4)
            rows = []
            for t in sorted(r.deposit_transactions, key=lambda x: x.time or "")[:50]:
                rows.append([
                    t.time[:19] if t.time else "",
                    t.currency,
                    t.change,
                    t.reason,
                    t.transaction_id[:20] if t.transaction_id else "—",
                    t.source_sheet if t.source_sheet else "deposit&withdraw&transfer_1",
                ])
            self._add_table(["Czas", "Waluta", "Zmiana", "Powód", "TxID", "Źródło"], rows, max_rows=50)
            if len(r.deposit_transactions) > 50:
                self._add_paragraph(f"... i {len(r.deposit_transactions)-50} więcej (pełna lista w JSON).", color=RGBColor(0x80, 0x80, 0x80))

        # Withdrawal — max 50
        if r.withdrawal_transactions:
            self._add_paragraph("")
            self._add_heading("Wypłaty (pierwsze 50)", level=4)
            rows = []
            for t in sorted(r.withdrawal_transactions, key=lambda x: x.time or "")[:50]:
                rows.append([
                    t.time[:19] if t.time else "",
                    t.currency,
                    t.change,
                    t.reason,
                    t.transaction_id[:20] if t.transaction_id else "—",
                    t.source_sheet if t.source_sheet else "deposit&withdraw&transfer_1",
                ])
            self._add_table(["Czas", "Waluta", "Zmiana", "Powód", "TxID", "Źródło"], rows, max_rows=50)
            if len(r.withdrawal_transactions) > 50:
                self._add_paragraph(f"... i {len(r.withdrawal_transactions)-50} więcej (pełna lista w JSON).", color=RGBColor(0x80, 0x80, 0x80))

        # Spot/Trade — max 50
        if r.spot_transactions:
            self._add_paragraph("")
            self._add_heading("Trade/Spot (pierwsze 50)", level=4)
            rows = []
            for t in sorted(r.spot_transactions, key=lambda x: x.time or "")[:50]:
                rows.append([
                    t.time[:19] if t.time else "",
                    t.currency,
                    t.change,
                    t.reason,
                    t.transaction_id[:20] if t.transaction_id else "—",
                    t.source_sheet if t.source_sheet else "trade",
                ])
            self._add_table(["Czas", "Waluta", "Zmiana", "Powód", "Order ID", "Źródło"], rows, max_rows=50)
            if len(r.spot_transactions) > 50:
                self._add_paragraph(f"... i {len(r.spot_transactions)-50} więcej (pełna lista w JSON).", color=RGBColor(0x80, 0x80, 0x80))

        # Bilans per waluta
        all_txns = r.deposit_transactions + r.withdrawal_transactions + r.spot_transactions + r.funding_transactions
        from collections import defaultdict
        per_curr = defaultdict(lambda: {"in": 0.0, "out": 0.0, "count": 0})
        for t in all_txns:
            curr = t.currency.upper() if t.currency else "UNKNOWN"
            try:
                chg = float(t.change) if t.change else 0.0
            except:
                chg = 0.0
            per_curr[curr]["count"] += 1
            if chg > 0: per_curr[curr]["in"] += chg
            elif chg < 0: per_curr[curr]["out"] += abs(chg)

        if per_curr:
            self._add_paragraph("")
            self._add_heading("Bilans per waluta", level=4)
            rows = []
            for curr, data in sorted(per_curr.items()):
                net = data["in"] - data["out"]
                rows.append([
                    curr,
                    f"+{data['in']:.8f}".rstrip('0').rstrip('.'),
                    f"-{data['out']:.8f}".rstrip('0').rstrip('.'),
                    f"{net:+.8f}".rstrip('0').rstrip('.'),
                    str(data["count"]),
                ])
            self._add_table(["Waluta", "Przychody", "Rozchody", "Netto", "Liczba trans."], rows, max_rows=50)
